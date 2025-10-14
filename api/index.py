"""
Vercel serverless function entry point for the transcript formatter.
"""

import os
import io
from http.server import BaseHTTPRequestHandler
import json
import anthropic
from docx import Document
from docx.shared import Pt
import cgi
import re

class handler(BaseHTTPRequestHandler):
    def do_GET(self):
        """Handle GET requests."""
        if self.path == '/' or self.path == '':
            self.send_response(200)
            self.send_header('Content-Type', 'text/html; charset=utf-8')
            self.end_headers()
            
            html_content = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Transcript Formatter - ORU</title>
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <style>
        :root {
            --oru-blue: #003366;
            --oru-gold: #FFD700;
            --oru-light-blue: #4A90E2;
            --oru-dark-blue: #002244;
            --oru-gray: #F5F5F5;
            --oru-white: #FFFFFF;
            --success-green: #28a745;
            --error-red: #dc3545;
            --warning-orange: #fd7e14;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, var(--oru-blue) 0%, var(--oru-dark-blue) 100%);
            min-height: 100vh;
            color: var(--oru-white);
        }

        .container {
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
        }

        .header {
            text-align: center;
            margin-bottom: 40px;
            padding: 40px 0;
        }

        .header h1 {
            font-size: 3rem;
            font-weight: 700;
            margin-bottom: 10px;
            background: linear-gradient(45deg, var(--oru-gold), #FFF8DC);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }

        .header p {
            font-size: 1.2rem;
            opacity: 0.9;
            margin-bottom: 20px;
        }

        .oru-logo {
            font-size: 1.1rem;
            color: var(--oru-gold);
            font-weight: 600;
        }

        .main-card {
            background: var(--oru-white);
            border-radius: 20px;
            padding: 40px;
            box-shadow: 0 20px 40px rgba(0, 0, 0, 0.1);
            color: var(--oru-blue);
            margin-bottom: 30px;
        }

        .upload-section {
            text-align: center;
            margin-bottom: 30px;
        }

        .upload-area {
            border: 3px dashed var(--oru-light-blue);
            border-radius: 15px;
            padding: 60px 20px;
            background: linear-gradient(135deg, #f8f9ff 0%, #e8f2ff 100%);
            transition: all 0.3s ease;
            cursor: pointer;
            position: relative;
            overflow: hidden;
        }

        .upload-area:hover {
            border-color: var(--oru-blue);
            background: linear-gradient(135deg, #e8f2ff 0%, #d0e7ff 100%);
            transform: translateY(-2px);
        }

        .upload-area.dragover {
            border-color: var(--oru-gold);
            background: linear-gradient(135deg, #fff8dc 0%, #ffeaa7 100%);
        }

        .upload-icon {
            font-size: 4rem;
            color: var(--oru-light-blue);
            margin-bottom: 20px;
        }

        .upload-text {
            font-size: 1.3rem;
            font-weight: 600;
            color: var(--oru-blue);
            margin-bottom: 10px;
        }

        .upload-subtext {
            color: #666;
            font-size: 1rem;
        }

        .file-input {
            display: none;
        }

        .options-section {
            margin: 30px 0;
            padding: 20px;
            background: var(--oru-gray);
            border-radius: 10px;
        }

        .process-btn {
            background: linear-gradient(135deg, var(--oru-blue) 0%, var(--oru-light-blue) 100%);
            color: var(--oru-white);
            border: none;
            padding: 15px 40px;
            font-size: 1.1rem;
            font-weight: 600;
            border-radius: 50px;
            cursor: pointer;
            transition: all 0.3s ease;
            box-shadow: 0 5px 15px rgba(0, 51, 102, 0.3);
            display: none;
        }

        .process-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 25px rgba(0, 51, 102, 0.4);
        }

        .process-btn:disabled {
            opacity: 0.6;
            cursor: not-allowed;
            transform: none;
        }

        .progress-section {
            display: none;
            text-align: center;
            margin: 30px 0;
        }

        .progress-bar {
            width: 100%;
            height: 8px;
            background: #e0e0e0;
            border-radius: 4px;
            overflow: hidden;
            margin: 20px 0;
        }

        .progress-fill {
            height: 100%;
            background: linear-gradient(90deg, var(--oru-blue), var(--oru-light-blue));
            width: 0%;
            transition: width 0.3s ease;
        }

        .result-section {
            display: none;
            margin-top: 30px;
        }

        .result-card {
            background: linear-gradient(135deg, var(--success-green) 0%, #20c997 100%);
            color: white;
            padding: 25px;
            border-radius: 15px;
            margin-bottom: 20px;
        }

        .result-card h3 {
            margin-bottom: 15px;
            font-size: 1.3rem;
        }

        .download-btn {
            background: var(--oru-gold);
            color: var(--oru-blue);
            border: none;
            padding: 12px 30px;
            font-size: 1rem;
            font-weight: 600;
            border-radius: 25px;
            cursor: pointer;
            transition: all 0.3s ease;
            text-decoration: none;
            display: inline-block;
        }

        .download-btn:hover {
            background: #e6c200;
            transform: translateY(-1px);
        }

        .error-message {
            background: var(--error-red);
            color: white;
            padding: 15px;
            border-radius: 10px;
            margin: 20px 0;
            display: none;
        }

        .feature-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin-top: 40px;
        }

        .feature-card {
            background: linear-gradient(135deg, var(--oru-white) 0%, var(--oru-gray) 100%);
            padding: 25px;
            border-radius: 15px;
            text-align: center;
            border: 2px solid transparent;
            transition: all 0.3s ease;
        }

        .feature-card:hover {
            border-color: var(--oru-gold);
            transform: translateY(-5px);
        }

        .feature-icon {
            font-size: 2.5rem;
            color: var(--oru-blue);
            margin-bottom: 15px;
        }

        .feature-title {
            font-size: 1.2rem;
            font-weight: 600;
            color: var(--oru-blue);
            margin-bottom: 10px;
        }

        .feature-desc {
            color: #666;
            font-size: 0.95rem;
        }

        .footer {
            text-align: center;
            margin-top: 50px;
            padding: 30px;
            border-top: 1px solid rgba(255, 255, 255, 0.1);
        }

        .footer p {
            opacity: 0.8;
            font-size: 0.9rem;
        }

        @media (max-width: 768px) {
            .container {
                padding: 10px;
            }
            
            .header h1 {
                font-size: 2rem;
            }
            
            .main-card {
                padding: 20px;
            }
            
            .upload-area {
                padding: 40px 15px;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1><i class="fas fa-file-alt"></i> Transcript Formatter</h1>
            <p>AI-Powered Transcript Formatting</p>
            <div class="oru-logo">
                <i class="fas fa-university"></i> Oral Roberts University
            </div>
        </div>

        <div class="main-card">
            <div class="upload-section">
                <div class="upload-area" id="uploadArea">
                    <div class="upload-icon">
                        <i class="fas fa-cloud-upload-alt"></i>
                    </div>
                    <div class="upload-text">Drop your transcript file here</div>
                    <div class="upload-subtext">or click to browse (.txt, .docx files)</div>
                    <input type="file" id="fileInput" class="file-input" accept=".txt,.docx">
                </div>
            </div>

            <div class="options-section">
                <p style="font-size: 1rem; color: #666; text-align: center; padding: 20px;">
                    <i class="fas fa-robot"></i> 
                    <strong>AI-Powered Formatting</strong><br>
                    Intelligent formatting with speaker detection and Scripture reference highlighting.
                </p>
            </div>

            <div style="text-align: center;">
                <button class="process-btn" id="processBtn">
                    <i class="fas fa-magic"></i> Format Transcript
                </button>
            </div>

            <div class="progress-section" id="progressSection">
                <h3><i class="fas fa-cog fa-spin"></i> Processing your transcript...</h3>
                <div class="progress-bar">
                    <div class="progress-fill" id="progressFill"></div>
                </div>
                <p id="progressText">Preparing transcript...</p>
            </div>

            <div class="error-message" id="errorMessage"></div>

            <div class="result-section" id="resultSection">
                <div class="result-card">
                    <h3><i class="fas fa-check-circle"></i> Transcript Formatted Successfully!</h3>
                    <p id="resultText">Your transcript has been processed and is ready for download.</p>
                    <div style="margin-top: 15px;">
                        <button class="download-btn" id="downloadBtn">
                            <i class="fas fa-download"></i> Download Formatted Document
                        </button>
                    </div>
                </div>
            </div>
        </div>

        <div class="feature-grid">
            <div class="feature-card">
                <div class="feature-icon">
                    <i class="fas fa-robot"></i>
                </div>
                <div class="feature-title">AI-Powered</div>
                <div class="feature-desc">Uses advanced AI for intelligent transcript formatting and speaker detection</div>
            </div>
            
            <div class="feature-card">
                <div class="feature-icon">
                    <i class="fas fa-bible"></i>
                </div>
                <div class="feature-title">Scripture Detection</div>
                <div class="feature-desc">Automatically detects and formats Scripture references in any format</div>
            </div>
            
            <div class="feature-card">
                <div class="feature-icon">
                    <i class="fas fa-users"></i>
                </div>
                <div class="feature-title">Speaker Formatting</div>
                <div class="feature-desc">Intelligently identifies and formats speaker names throughout the transcript</div>
            </div>
            
            <div class="feature-card">
                <div class="feature-icon">
                    <i class="fas fa-file-word"></i>
                </div>
                <div class="feature-title">Professional Output</div>
                <div class="feature-desc">Generates clean, professional Word documents ready for distribution</div>
            </div>
        </div>

        <div class="footer">
            <p>&copy; 2025 Oral Roberts University - Transcript Formatter</p>
            <p>Powered by AI and built with ❤️</p>
        </div>
    </div>

    <script>
        const uploadArea = document.getElementById('uploadArea');
        const fileInput = document.getElementById('fileInput');
        const processBtn = document.getElementById('processBtn');
        const progressSection = document.getElementById('progressSection');
        const progressFill = document.getElementById('progressFill');
        const progressText = document.getElementById('progressText');
        const resultSection = document.getElementById('resultSection');
        const downloadBtn = document.getElementById('downloadBtn');
        const errorMessage = document.getElementById('errorMessage');
        const resultText = document.getElementById('resultText');

        let selectedFile = null;
        let downloadBlob = null;

        // Upload area click handler
        uploadArea.addEventListener('click', () => {
            fileInput.click();
        });

        // Drag and drop handlers
        uploadArea.addEventListener('dragover', (e) => {
            e.preventDefault();
            uploadArea.classList.add('dragover');
        });

        uploadArea.addEventListener('dragleave', () => {
            uploadArea.classList.remove('dragover');
        });

        uploadArea.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            const files = e.dataTransfer.files;
            if (files.length > 0) {
                handleFileSelect(files[0]);
            }
        });

        // File input change handler
        fileInput.addEventListener('change', (e) => {
            if (e.target.files.length > 0) {
                handleFileSelect(e.target.files[0]);
            }
        });

        // Process button click handler
        processBtn.addEventListener('click', () => {
            if (selectedFile) {
                processFile();
            }
        });

        // Download button click handler
        downloadBtn.addEventListener('click', () => {
            if (downloadBlob) {
                const url = window.URL.createObjectURL(downloadBlob);
                const a = document.createElement('a');
                a.href = url;
                a.download = selectedFile.name.replace(/\.[^/.]+$/, '') + '_formatted.docx';
                document.body.appendChild(a);
                a.click();
                document.body.removeChild(a);
                window.URL.revokeObjectURL(url);
            }
        });

        function handleFileSelect(file) {
            const allowedTypes = ['text/plain', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'];
            
            if (!allowedTypes.includes(file.type) && !file.name.toLowerCase().endsWith('.txt') && !file.name.toLowerCase().endsWith('.docx')) {
                showError('Please select a .txt or .docx file.');
                return;
            }

            selectedFile = file;
            
            // Update upload area
            uploadArea.innerHTML = '<div class="upload-icon" style="color: var(--success-green);"><i class="fas fa-file-check"></i></div><div class="upload-text" style="color: var(--success-green);">File Selected: ' + file.name + '</div><div class="upload-subtext">Click "Format Transcript" to process</div>';
            
            processBtn.style.display = 'inline-block';
            hideError();
            hideResult();
        }

        function processFile() {
            if (!selectedFile) return;

            const formData = new FormData();
            formData.append('file', selectedFile);

            // Show progress
            showProgress();
            hideError();
            hideResult();

            // Simulate progress updates
            let progress = 0;
            const progressInterval = setInterval(() => {
                progress += Math.random() * 15;
                if (progress > 90) progress = 90;
                updateProgress(progress, 'Processing transcript...');
            }, 500);

            fetch('/api/upload', {
                method: 'POST',
                body: formData
            })
            .then(response => {
                if (!response.ok) {
                    throw new Error('Processing failed');
                }
                return response.blob();
            })
            .then(blob => {
                clearInterval(progressInterval);
                updateProgress(100, 'Complete!');
                downloadBlob = blob;
                
                setTimeout(() => {
                    hideProgress();
                    showResult();
                }, 1000);
            })
            .catch(error => {
                clearInterval(progressInterval);
                hideProgress();
                showError('Processing error: ' + error.message);
            });
        }

        function showProgress() {
            progressSection.style.display = 'block';
            processBtn.disabled = true;
        }

        function hideProgress() {
            progressSection.style.display = 'none';
            processBtn.disabled = false;
        }

        function updateProgress(percent, text) {
            progressFill.style.width = percent + '%';
            progressText.textContent = text;
        }

        function showResult() {
            resultSection.style.display = 'block';
            resultText.textContent = 'Your document has been formatted successfully using AI!';
        }

        function hideResult() {
            resultSection.style.display = 'none';
        }

        function showError(message) {
            errorMessage.textContent = message;
            errorMessage.style.display = 'block';
        }

        function hideError() {
            errorMessage.style.display = 'none';
        }
    </script>
</body>
</html>'''
            
            self.wfile.write(html_content.encode('utf-8'))
            
        elif self.path == '/health':
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            
            try:
                # Test imports
                import anthropic
                from docx import Document
                response = {'status': 'healthy', 'service': 'transcript-formatter'}
            except ImportError as e:
                response = {'status': 'unhealthy', 'error': str(e)}
            
            self.wfile.write(json.dumps(response).encode())
            
        else:
            self.send_response(404)
            self.end_headers()
            self.wfile.write(b'Not found')
    
    def do_POST(self):
        """Handle POST requests."""
        if self.path == '/api/upload':
            try:
                # Parse the form data
                content_type = self.headers['content-type']
                if not content_type or 'multipart/form-data' not in content_type:
                    self.send_error(400, 'Invalid content type')
                    return
                
                # Get boundary
                ctype, pdict = cgi.parse_header(content_type)
                pdict['boundary'] = pdict['boundary'].encode()
                
                # Parse multipart form data
                content_length = int(self.headers['Content-Length'])
                form = cgi.FieldStorage(
                    fp=self.rfile,
                    headers=self.headers,
                    environ={'REQUEST_METHOD': 'POST'}
                )
                
                # Get the file
                if 'file' not in form:
                    self.send_error(400, 'No file uploaded')
                    return
                
                file_item = form['file']
                if not file_item.filename:
                    self.send_error(400, 'No file selected')
                    return
                
                # Check file extension
                filename = file_item.filename
                if not (filename.endswith('.txt') or filename.endswith('.docx')):
                    self.send_error(400, 'Invalid file type')
                    return
                
                # Read file content
                file_content = file_item.file.read()
                
                # Extract text based on file type
                if filename.endswith('.docx'):
                    doc = Document(io.BytesIO(file_content))
                    content = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    content = file_content.decode('utf-8')
                
                # Format with Claude
                api_key = os.environ.get('ANTHROPIC_API_KEY')
                if not api_key:
                    self.send_error(500, 'API key not configured')
                    return
                
                client = anthropic.Anthropic(api_key=api_key)
                
                prompt = """You are a professional transcript editor. Transform this raw transcript into a well-formatted, readable document.

Please follow these guidelines:
1. Create clear paragraph breaks for better readability
2. Add proper punctuation and capitalization
3. Identify and format speaker names (if present) consistently
4. Remove filler words and false starts while preserving meaning
5. Organize the content into logical sections
6. Clean up any formatting issues or artifacts

Raw Transcript:
{transcript}

Return the formatted transcript ready for a professional document."""
                
                response = client.messages.create(
                    model="claude-3-haiku-20240307",
                    max_tokens=4000,
                    temperature=0.3,
                    messages=[
                        {
                            "role": "user",
                            "content": prompt.format(transcript=content[:10000])
                        }
                    ]
                )
                formatted_text = response.content[0].text
                
                # Create Word document
                doc = Document()
                doc.add_heading('Formatted Transcript', 0)
                
                # Process formatted text
                paragraphs = formatted_text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        if ':' in para_text and len(para_text.split(':')[0]) < 50:
                            p = doc.add_paragraph()
                            speaker_parts = para_text.split(':', 1)
                            speaker_run = p.add_run(speaker_parts[0] + ':')
                            speaker_run.bold = True
                            if len(speaker_parts) > 1:
                                p.add_run(' ' + speaker_parts[1].strip())
                        else:
                            doc.add_paragraph(para_text.strip())
                
                # Save to bytes
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                # Send response
                self.send_response(200)
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Disposition', 'attachment; filename=formatted_transcript.docx')
                self.end_headers()
                self.wfile.write(docx_buffer.getvalue())
                
            except Exception as e:
                self.send_error(500, str(e))
        else:
            self.send_error(404, 'Not found')