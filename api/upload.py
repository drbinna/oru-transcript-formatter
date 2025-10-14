"""
Upload handler for Vercel - working file upload processor.
"""

from http.server import BaseHTTPRequestHandler
import json
import os
import io
import cgi
import anthropic
from docx import Document

class handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
    
    def do_POST(self):
        try:
            # Get content type and length
            content_type = self.headers.get('content-type', '')
            content_length = int(self.headers.get('content-length', 0))
            
            if content_length == 0:
                self.send_error(400, 'No data received')
                return
            
            # Use cgi.FieldStorage for proper multipart parsing
            environ = {
                'REQUEST_METHOD': 'POST',
                'CONTENT_TYPE': content_type,
                'CONTENT_LENGTH': str(content_length),
            }
            
            # Parse the multipart form data
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ=environ
            )
            
            # Get the uploaded file
            if 'file' not in form:
                self.send_error(400, 'No file field found')
                return
            
            file_item = form['file']
            if not hasattr(file_item, 'filename') or not file_item.filename:
                self.send_error(400, 'No file uploaded')
                return
            
            filename = file_item.filename
            if not (filename.lower().endswith('.txt') or filename.lower().endswith('.docx')):
                self.send_error(400, 'Only .txt and .docx files supported')
                return
            
            # Read file content
            file_content = file_item.file.read()
            
            # Extract text
            if filename.lower().endswith('.docx'):
                try:
                    doc = Document(io.BytesIO(file_content))
                    text_content = '\\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                except Exception as e:
                    self.send_error(400, f'Error reading Word document: {str(e)}')
                    return
            else:
                try:
                    text_content = file_content.decode('utf-8', errors='ignore')
                except Exception as e:
                    self.send_error(400, f'Error reading text file: {str(e)}')
                    return
            
            if not text_content.strip():
                self.send_error(400, 'File is empty')
                return
            
            # Format with Claude
            api_key = os.environ.get('ANTHROPIC_API_KEY')
            if not api_key:
                self.send_error(500, 'API key not configured')
                return
            
            try:
                client = anthropic.Anthropic(api_key=api_key)
                
                prompt = f\"\"\"You are a professional transcript editor. Transform this raw transcript into a well-formatted, readable document.

Please follow these guidelines:
1. Create clear paragraph breaks for better readability
2. Add proper punctuation and capitalization
3. Identify and format speaker names (if present) consistently
4. Remove filler words and false starts while preserving meaning
5. Organize the content into logical sections
6. Clean up any formatting issues or artifacts

Raw Transcript:
{text_content[:8000]}

Return the formatted transcript ready for a professional document.\"\"\"
                
                response = client.messages.create(
                    model="claude-3-haiku-20240307",
                    max_tokens=4000,
                    temperature=0.3,
                    messages=[{"role": "user", "content": prompt}]
                )
                
                formatted_text = response.content[0].text
                
            except Exception as e:
                self.send_error(500, f'AI processing failed: {str(e)}')
                return
            
            # Create Word document
            try:
                doc = Document()
                doc.add_heading('Formatted Transcript', 0)
                
                # Add formatted content
                paragraphs = formatted_text.split('\\n\\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # Check for speaker format
                        if ':' in para_text and len(para_text.split(':')[0]) < 50:
                            # Speaker line
                            p = doc.add_paragraph()
                            speaker_parts = para_text.split(':', 1)
                            speaker_run = p.add_run(speaker_parts[0] + ':')
                            speaker_run.bold = True
                            if len(speaker_parts) > 1:
                                p.add_run(' ' + speaker_parts[1].strip())
                        else:
                            # Regular paragraph
                            doc.add_paragraph(para_text.strip())
                
                # Save to buffer
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                docx_data = docx_buffer.getvalue()
                
                # Send response
                self.send_response(200)
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Disposition', f'attachment; filename="formatted_{filename.replace(".txt", ".docx")}"')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.send_header('Content-Length', str(len(docx_data)))
                self.end_headers()
                
                self.wfile.write(docx_data)
                
            except Exception as e:
                self.send_error(500, f'Document creation failed: {str(e)}')
                return
                
        except Exception as e:
            self.send_error(500, f'Server error: {str(e)}')
            return