"""
Flask web application for the transcript formatter.
Provides a modern HTML interface for uploading and formatting transcripts.
"""

import os
import tempfile
import traceback
import logging
from pathlib import Path
from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Core dependencies for AI formatting
import anthropic
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__, static_folder='static', static_url_path='/static')
CORS(app)  # Enable CORS for all routes
app.secret_key = os.environ.get('FLASK_SECRET_KEY', os.urandom(24).hex())  # Secure secret key
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Request logging (simplified)
@app.before_request
def log_request_info():
    """Log all incoming requests for debugging."""
    try:
        logger.info(f"Request: {request.method} {request.path}")
    except Exception as e:
        logger.error(f"Error in before_request: {e}")

# Global error handler to ensure JSON responses
@app.errorhandler(Exception)
def handle_exception(e):
    """Handle all unhandled exceptions and return JSON."""
    app.logger.error(f"Unhandled exception: {str(e)}", exc_info=True)
    response = jsonify({
        'success': False,
        'error': str(e)
    })
    response.headers['Content-Type'] = 'application/json'
    return response, 500

@app.errorhandler(500)
def handle_500_error(e):
    """Handle 500 errors and return JSON."""
    app.logger.error(f"500 error: {str(e)}")
    response = jsonify({
        'success': False,
        'error': 'Internal server error'
    })
    response.headers['Content-Type'] = 'application/json'
    return response, 500

# Add more specific error handlers
@app.errorhandler(404)
def handle_404_error(e):
    """Handle 404 errors and return JSON for API endpoints."""
    if request.path.startswith('/upload') or request.path.startswith('/download') or request.path.startswith('/debug') or request.path.startswith('/health'):
        response = jsonify({
            'success': False,
            'error': 'Endpoint not found'
        })
        response.headers['Content-Type'] = 'application/json'
        return response, 404
    # For other paths, return normal 404
    return render_template('index.html'), 404

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
ALLOWED_EXTENSIONS = {'txt', 'docx'}

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Claude AI formatting functionality
def format_with_claude_inline(transcript_text):
    """Format transcript using Claude AI - inline implementation."""
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    
    if not api_key:
        raise ValueError("Anthropic API key not found. Please set ANTHROPIC_API_KEY environment variable.")
    
    client = anthropic.Anthropic(api_key=api_key)
    
    system_prompt = """You are a professional transcript formatter that converts raw AI-generated transcripts into polished, publication-ready documents. Output clean text WITHOUT any asterisks, underscores, or markdown symbols. The Word document exporter will handle all formatting. Document body will use Times New Roman size 12, while the title will be centered, bold, underlined in Gotham size 20. A branded template with pre-configured header and footers will be used for all documents.

<divider_line_rules>

## DIVIDER LINE IMPLEMENTATION

Use this exact divider line (80 dashes):
────────────────────────────────────────────────────────────────────────────────

### DIVIDER PLACEMENT RULES:

1. **Title Section**
   - Insert divider ABOVE the title
   - Insert divider BELOW the title
   - Add blank line before and after each divider

2. **Speaker Changes**
   - Insert divider ABOVE each new speaker
   - Exception: Skip if same speaker continues (use "Speaker (continued):" instead)
   - Exception: Skip for minor interjections within same topic

3. **Major Theme Shifts**
   - Insert divider when transitioning between:
     * Narration → Song
     * Testimony → Sermon
     * Prayer → Announcement
     * Teaching → Music
   - Detect via keywords: ♪, Song, Music, Prayer, Teaching, Scripture

4. **Numbered Sections**
   - Insert divider ABOVE numbered headings (1. Section Title)
   - Do not add divider below the heading

5. **Song/Music Sections**
   - Insert divider ABOVE the start of lyrics
   - Insert divider BELOW the end of lyrics when returning to prose

6. **Return to Main Speaker**
   - Insert divider when returning to main speaker after interlude

7. **Closing Credits**
   - Insert divider ABOVE credits/license/outro section
   - Detect: "Presented by", "This has been", "Copyright", "License"

### FORMATTING SAFEGUARDS:
- **NEVER place two or more consecutive dividers** - if multiple dividers would occur, keep only ONE
- **NEVER duplicate dividers** - if a divider already exists after title, do not add another
- Never begin or end the document with a divider (exception: divider before title is allowed)
- Always add EXACTLY ONE blank line above and ONE blank line below each divider
- Maintain consistent 80-dash width for all dividers
- Do not use dividers for minor transitions within same speaker's content
- **Consolidation rule:** If logic suggests multiple dividers in sequence, consolidate to single divider
- **Title exception:** After the title block divider, skip the next speaker divider if they're the same person

</divider_line_rules>

<formatting_rules>

## 1. TITLE EXTRACTION AND PLACEMENT
<title_rules>
- Extract the main title from context or filename
- Place at the very top of the document on its own line
- CRITICAL: Title MUST be formatted as CENTERED, BOLD, UNDERLINED, Gotham font size 20 in final Word document
- The title should be the first line of content (after opening divider)
- Add blank line after title
- Example: Living in the Last Days
- NOTE: Output plain text - the Word document exporter will automatically apply:
  * CENTER alignment
  * BOLD formatting
  * UNDERLINE formatting  
  * Gotham font, size 20pt (Times New Roman fallback)
</title_rules>

## 2. SPEAKER FORMATTING
<speaker_rules>
- Format all speakers as: Speaker Name: (with colon and space)
- Remove ALL timestamps: [Laura Lacy] 09:03:24 → Laura Lacy:
- For continuing speakers after interruption: Speaker Name (continued):
- Add blank line between different speakers
- Group same speaker's consecutive lines into coherent paragraphs (3-6 sentences)
- Example: Dr. Billy Wilson: Welcome to World Impact.
</speaker_rules>

## 3. SCRIPTURE REFERENCES
<scripture_rules>
- Format Bible references as: Book Chapter:Verse or Book Chapter:Verse--Verse
- Recognize patterns like:
  - "1 John chapter 2, verse 18" → 1 John 2:18
  - "2 Timothy 3, verse 1 through 5" → 2 Timothy 3:1--5
  - "Mark chapter 13, verse 13" → Mark 13:13
- Keep the actual Scripture quotes in quotes
- Example: Hebrews 5:14 says, "But strong meat belongs to them..."
</scripture_rules>

## 4. SECTION NUMBERING
<section_rules>
- Detect numbered sections in the content
- Format section headers as: 1. Section Title Here
- Recognize transitional phrases:
  - "The first is..." → 1. [Topic]
  - "The second thing..." → 2. [Topic]
  - "And finally..." / "And most importantly..." → 5. [Topic]
- Add blank line before and after section headers
- Example: 1. A Counterculture Mindset
</section_rules>

## 5. SPECIAL TEXT FORMATTING
<special_formatting>
- Show/Program names: Keep in quotes: "World Impact"
- Organizations: Keep as is: ORU, Oral Roberts University
- Websites: Keep as is: worldimpact.tv
- Song titles: Keep in quotes: "Give Me Jesus"
- Quoted speech: Keep in quotes: "Hey, bring that boy to me"
</special_formatting>

## 6. CHARACTER ENCODING FIXES
<encoding_rules>
Replace broken characters:
- â™ª → ♪
- â€™ → '
- â€œ → "
- â€ → "
- â€" → --
- Iâ€™m → I'm
- youâ€™re → you're
</encoding_rules>

## 7. CONTENT CLEANUP
<cleanup_rules>
- Remove "..." at beginning/end of lines
- Fix stutters: "we know the--we need" → "we need"
- Remove excessive filler words (um, uh, like) but preserve natural speech
- Remove standalone music symbol lines (lines with only â™ª or ♪)
- Clean up multiple spaces → single space
- Remove timestamps completely
</cleanup_rules>

## 8. PARAGRAPH STRUCTURE
<paragraph_rules>
- Create natural 3-6 sentence paragraphs
- Merge fragmented sentences from same speaker into flowing text
- Keep related thoughts together
- Add blank line between different speakers
- Use Speaker (continued): when same speaker resumes after interruption
- Preserve paragraph breaks for topic changes
</paragraph_rules>

## 9. MUSIC FORMATTING
<music_rules>
- Format song lyrics as: ♪ Lyric line here. ♪
- Keep music symbols clean: single ♪ at start and end
- Group related lyrics together
- Add blank line before/after music sections
</music_rules>

## 10. PROFESSIONAL POLISH
<polish_rules>
- Proper capitalization for names and titles
- Consistent punctuation
- Proper spacing around colons, periods, commas
- No double spaces
- No orphaned punctuation
- Clean line breaks
</polish_rules>

</formatting_rules>

<footer_requirement>

## MANDATORY FOOTER

The Word document exporter will automatically add this footer to every page:

"Oral Roberts University Presents: World Impact with Dr. Billy Wilson"

Footer Rules:
- The Word document exporter will automatically add this footer to EVERY PAGE
- Do NOT include the footer text in your transcript output
- The footer will appear at the bottom of every page in the final Word document
- Footer text: "Oral Roberts University Presents: World Impact with Dr. Billy Wilson"
- Footer styling: Arial 10pt, light gray, centered
- This footer is MANDATORY for all transcripts and handled automatically

</footer_requirement>

<example_transformation>
INPUT:
```
... â™ªâ™ªâ™ª Dr. Billy Wilson: Welcome to "World Impact." Today we are in Krakow, Poland...
Well, I wanna talk about five things I believe you need in your life in order to live successfully in the last days. The first is a counterculture mindset...
Male Announcer: Ever since Jesus said...
♪ Give me Jesus ♪
Billy: Today on World Impact...
```

OUTPUT:
```
────────────────────────────────────────────────────────────────────────────────

Living in the Last Days

────────────────────────────────────────────────────────────────────────────────

Dr. Billy Wilson: Welcome to "World Impact." I'm Billy Wilson. And today we are in Krakow, Poland...

Billy (continued): Well, I wanna talk about five things I believe you need in your life in order to live successfully in the last days.

────────────────────────────────────────────────────────────────────────────────

1. A Counterculture Mindset

We live in a culture filled with dishonor and impurity...

────────────────────────────────────────────────────────────────────────────────

Male Announcer: Ever since Jesus said...

────────────────────────────────────────────────────────────────────────────────

♪ Give me Jesus ♪

────────────────────────────────────────────────────────────────────────────────

Billy: Today on World Impact...

[... rest of transcript content ...]
```
</example_transformation>

<critical_notes>
- DO NOT use asterisks, underscores, or any markdown formatting
- Output plain clean text only
- TITLE MUST be first line after opening divider - will be formatted as CENTERED, BOLD, UNDERLINED, Gotham 20pt
- ALWAYS use 80-dash divider lines: ────────────────────────────────────────────────────────────────────────────────
- ONLY ONE DIVIDER between sections - NEVER consecutive dividers
- If two dividers would appear in sequence, use only ONE
- Do not add divider after title if speaker immediately follows
- ALWAYS add EXACTLY one blank line before and after each divider
- NEVER begin or end document with a divider (except before title)
- Format speaker names with colon: Name:
- Format Scripture references cleanly: Book Chapter:Verse
- Number sections when content indicates them
- Fix all encoding issues
- Create coherent paragraphs (3-6 sentences)
- Add Speaker (continued): for interrupted dialogue
- Do NOT include footer text in output - Word exporter adds it automatically to every page
- NO ASTERISKS OR MARKDOWN - plain text output only
- REMEMBER: Word exporter handles all formatting - just output clean text
</critical_notes>

Now format the transcript:"""
    
    try:
        # Use Claude Sonnet 4.5 - the latest model
        logger.info("Calling Claude Sonnet 4.5 API...")
        response = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=8192,
            temperature=0.1,
            system=system_prompt,
            messages=[
                {
                    "role": "user",
                    "content": f"Please format this transcript:\n\n{transcript_text}"
                }
            ]
        )
        
        # Get the formatted text from the response
        formatted_text = response.content[0].text
        logger.info("Claude Sonnet 4.5 API call successful")
        
        return formatted_text
            
    except anthropic.APIError as e:
        logger.error(f"Claude API Error: Status={e.status_code if hasattr(e, 'status_code') else 'N/A'}, Message={str(e)}")
        raise RuntimeError(f"Claude API error: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error calling Claude: {type(e).__name__}: {str(e)}")
        raise RuntimeError(f"Claude API error: {str(e)}")

def create_word_document(formatted_text, title, output_path):
    """Create a professionally formatted Word document using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import re
        import os
        
        # ALWAYS use template - it's required for all documents
        template_path = os.path.join('static', 'template.docx')
        
        if os.path.exists(template_path):
            # Use the template document with pre-configured header and footer
            doc = Document(template_path)
            logger.info("Using required template document with pre-configured branding")
        else:
            # Fallback if template is missing
            doc = Document()
            logger.warning("TEMPLATE NOT FOUND - template.docx is required in static folder!")
        
        # Extract title from the formatted text (first non-divider line)
        text_lines = formatted_text.split('\n')
        document_title = title  # Default fallback
        title_line_index = -1
        
        # Find the first line that's not empty or a divider
        for i, line in enumerate(text_lines):
            line = line.strip()
            if line and not line.startswith('─'):  # Not empty and not a divider
                document_title = line
                title_line_index = i
                break
        
        # Remove title from the content if we found it
        if title_line_index >= 0:
            # Keep content after the title line
            formatted_text = '\n'.join(text_lines[title_line_index + 1:])
        
        # Add title (bold, centered, underlined, Gotham/Times New Roman 20)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_para.add_run(document_title)
        title_run.bold = True
        title_run.underline = True
        # Try Gotham first, fallback to Times New Roman
        try:
            title_run.font.name = 'Gotham'
        except:
            title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(20)
        
        # Add blank line after title
        doc.add_paragraph()
        
        # Process formatted text line by line
        lines = formatted_text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line
                doc.add_paragraph('')
                continue
            
            # Create paragraph
            p = doc.add_paragraph()
            
            # Check if this is a speaker name (ends with colon)
            if re.match(r'^[A-Za-z\s\(\)]+:$', line):
                # Speaker name - make it bold
                run = p.add_run(line)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
            # Check if this is a section header (starts with number and period)
            elif re.match(r'^\d+\.\s+.+', line):
                # Section header - make it bold
                run = p.add_run(line)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
            else:
                # Regular content - process for Scripture references and special formatting
                # Split by Scripture references pattern (Book Chapter:Verse)
                scripture_pattern = r'(\b(?:Genesis|Exodus|Leviticus|Numbers|Deuteronomy|Joshua|Judges|Ruth|1 Samuel|2 Samuel|1 Kings|2 Kings|1 Chronicles|2 Chronicles|Ezra|Nehemiah|Esther|Job|Psalm|Psalms|Proverbs|Ecclesiastes|Song of Songs|Isaiah|Jeremiah|Lamentations|Ezekiel|Daniel|Hosea|Joel|Amos|Obadiah|Jonah|Micah|Nahum|Habakkuk|Zephaniah|Haggai|Zechariah|Malachi|Matthew|Mark|Luke|John|Acts|Romans|1 Corinthians|2 Corinthians|Galatians|Ephesians|Philippians|Colossians|1 Thessalonians|2 Thessalonians|1 Timothy|2 Timothy|Titus|Philemon|Hebrews|James|1 Peter|2 Peter|1 John|2 John|3 John|Jude|Revelation)\s+\d+:\d+(?:--\d+)?)\b'
                
                parts = re.split(scripture_pattern, line)
                
                for part in parts:
                    if not part:
                        continue
                    
                    # Check if this part is a Scripture reference
                    if re.match(scripture_pattern, part):
                        # Scripture reference - make it bold
                        run = p.add_run(part)
                        run.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    else:
                        # Regular text
                        run = p.add_run(part)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
        
        # Add the gray footer that appears on every page
        section = doc.sections[0]
        footer = section.footer
        
        # Clear any existing footer content first
        for para in footer.paragraphs:
            para.clear()
        
        # Add the standard page footer
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add footer text
        footer_text = "Oral Roberts University Presents: World Impact with Dr. Billy Wilson"
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.name = 'Arial'
        footer_run.font.size = Pt(10)
        # Set light gray color
        try:
            footer_run.font.color.rgb = RGBColor(187, 187, 187)  # Light gray #BBBBBB
        except:
            pass  # Fall back to default color if RGBColor fails
        
        # Add Creative Commons license at the very end of the document
        doc.add_paragraph()  # Add blank line before license
        doc.add_paragraph()  # Add another blank line
        
        # Add the Creative Commons license text
        license_para = doc.add_paragraph()
        license_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        license_text = "This content available for use under a Creative Commons Attribution-NonCommercial license."
        license_run = license_para.add_run(license_text)
        license_run.font.name = 'Times New Roman'
        license_run.font.size = Pt(10)
        license_run.font.italic = True
        try:
            license_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
        except:
            pass
        
        # Save document
        doc.save(output_path)
        logger.info(f"Word document saved successfully: {output_path}")
        
    except ImportError:
        logger.warning("python-docx not available, saving as text file")
        # Fallback to text file if python-docx not available
        with open(output_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
            f.write(f"Title: {title}\n\n")
            f.write(formatted_text)

@app.route('/')
def index():
    """Main page with upload form."""
    return render_template('index.html')

@app.route('/upload', methods=['POST', 'OPTIONS'])
def upload_file():
    """Handle file upload and processing."""
    # Handle preflight OPTIONS request
    if request.method == 'OPTIONS':
        response = jsonify({'status': 'ok'})
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return response
    
    try:
        logger.info("=== UPLOAD REQUEST START ===")
        logger.info(f"Request method: {request.method}")
        logger.info(f"Request path: {request.path}")
        logger.info(f"Request content type: {request.content_type}")
        logger.info(f"Files in request: {list(request.files.keys()) if request.files else 'None'}")
        
        
        if 'file' not in request.files:
            logger.warning("No file in request")
            response = jsonify({'success': False, 'error': 'No file selected'})
            response.headers['Content-Type'] = 'application/json'
            return response, 400
        
        file = request.files['file']
        logger.info(f"File received: {file.filename}")
        
        if file.filename == '':
            logger.warning("Empty filename")
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        if file and allowed_file(file.filename):
            upload_path = None
            try:
                # Save uploaded file
                logger.info("=== FILE PROCESSING START ===")
                filename = secure_filename(file.filename)
                upload_path = os.path.join(UPLOAD_FOLDER, filename)
                logger.info(f"Saving file to: {upload_path}")
                
                try:
                    file.save(upload_path)
                    logger.info("File saved successfully")
                except Exception as save_error:
                    logger.error(f"File save failed: {save_error}")
                    response = jsonify({'success': False, 'error': f'File save failed: {str(save_error)}'})
                    response.headers['Content-Type'] = 'application/json'
                    return response, 500
                
                # Read file content
                logger.info("Reading file content")
                try:
                    with open(upload_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    logger.info(f"File content length: {len(content)} characters")
                except Exception as read_error:
                    logger.error(f"File read failed: {read_error}")
                    # Try different encodings
                    try:
                        with open(upload_path, 'r', encoding='latin-1') as f:
                            content = f.read()
                        logger.info(f"File read with latin-1 encoding, length: {len(content)} characters")
                    except Exception as read_error2:
                        logger.error(f"File read failed with latin-1: {read_error2}")
                        if upload_path and os.path.exists(upload_path):
                            os.remove(upload_path)
                        response = jsonify({'success': False, 'error': f'File read failed: {str(read_error)}'})
                        response.headers['Content-Type'] = 'application/json'
                        return response, 500
                
                # Format the transcript using AI
                logger.info("Starting AI formatting")
                try:
                    formatted_text = format_with_claude_inline(content)
                    formatter_used = 'Claude Sonnet 4.5'
                    logger.info("AI formatting completed successfully")
                except Exception as e:
                    logger.error(f"AI formatting failed: {str(e)}")
                    logger.error(f"AI formatting traceback: {traceback.format_exc()}")
                    # Clean up uploaded file on error
                    if upload_path and os.path.exists(upload_path):
                        os.remove(upload_path)
                    response = jsonify({'success': False, 'error': f'AI formatting failed: {str(e)}'})
                    response.headers['Content-Type'] = 'application/json'
                    return response, 500
                
                # Create output filename
                base_name = Path(filename).stem
                output_filename = f"{base_name}_formatted.docx"
                output_path = os.path.join(OUTPUT_FOLDER, output_filename)
                logger.info(f"Creating Word document: {output_path}")
                
                # Create Word document
                title = base_name.replace('_', ' ').replace('-', ' ')
                create_word_document(formatted_text, title, output_path)
                
                # Clean up uploaded file
                if upload_path and os.path.exists(upload_path):
                    os.remove(upload_path)
                
                logger.info("Upload processing completed successfully")
                response = jsonify({
                    'success': True,
                    'filename': output_filename,
                    'formatter': formatter_used,
                    'preview': formatted_text[:500] + '...' if len(formatted_text) > 500 else formatted_text
                })
                response.headers['Content-Type'] = 'application/json'
                return response
                
            except Exception as e:
                logger.error(f"Processing error: {str(e)}")
                logger.error(f"Processing traceback: {traceback.format_exc()}")
                # Clean up uploaded file on any error
                if upload_path and os.path.exists(upload_path):
                    os.remove(upload_path)
                return jsonify({'success': False, 'error': f'Processing failed: {str(e)}'}), 500
        
        logger.warning(f"Invalid file type: {file.filename}")
        return jsonify({'success': False, 'error': 'Invalid file type. Please upload .txt or .docx files.'}), 400
    
    except Exception as e:
        logger.error(f"Unhandled upload error: {str(e)}")
        logger.error(f"Unhandled upload traceback: {traceback.format_exc()}")
        # Catch all unhandled exceptions and return JSON
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Download processed file."""
    try:
        file_path = os.path.join(OUTPUT_FOLDER, secure_filename(filename))
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return jsonify({'success': False, 'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'success': False, 'error': f'Download failed: {str(e)}'}), 500

@app.route('/health')
def health_check():
    """Health check endpoint."""
    return jsonify({'status': 'healthy', 'service': 'transcript-formatter'})

@app.route('/debug')
def debug_env():
    """Minimal debug endpoint."""
    response = jsonify({
        'success': True,
        'message': 'Debug endpoint working',
        'version': 'minimal',
        'env_count': len(os.environ.keys())
    })
    response.headers['Content-Type'] = 'application/json'
    return response

@app.route('/test')
def test_endpoint():
    """Simple test endpoint to verify JSON responses."""
    try:
        return jsonify({
            'success': True,
            'message': 'Test successful',
            'timestamp': str(os.environ.get('REQUEST_ID', 'unknown'))
        })
    except Exception as e:
        return f"Error: {str(e)}", 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8081))
    app.run(debug=False, host='0.0.0.0', port=port)