"""
Fast transcript summarizer using Fireworks AI (OpenAI-compatible API).
Single API call, no chunking - processes the whole transcript at once.
"""

import os
import io
from openai import OpenAI, AuthenticationError, RateLimitError, APIStatusError, APIConnectionError, APITimeoutError
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from dotenv import load_dotenv

load_dotenv()

SUMMARIZE_PROMPT = """You are a professional transcript summarizer. Your job is to read a raw meeting or video transcript and produce a clean, readable summary document.

Structure your output EXACTLY like this:

TITLE: [A concise title for this transcript]

OVERVIEW:
[2-3 sentence summary of what this transcript is about]

KEY POINTS:
- [Key point 1]
- [Key point 2]
- [Key point 3]
- [Add more as needed, each on its own line starting with -]

MAIN DISCUSSION:
[3-5 paragraphs summarizing the main content in clear, professional prose. Capture the essential ideas, decisions, and discussion in a logical flow.]

ACTION ITEMS (if any):
- [Action item 1, or write "None identified" if there are no action items]

Keep the summary concise, professional, and easy to read. Remove filler words, repetition, and irrelevant chatter."""


FIREWORKS_BASE_URL = "https://api.fireworks.ai/inference/v1"
DEFAULT_MODEL = "accounts/fireworks/models/deepseek-v4-flash-0731"


class SummarizerError(Exception):
    """User-facing error with a safe, friendly message."""


def format_transcript(text: str, title: str = None) -> bytes:
    """Summarize a transcript with a single Fireworks call. Returns .docx bytes."""
    api_key = os.getenv('FIREWORKS_API_KEY')
    if not api_key:
        raise SummarizerError("The summarizer is not configured (missing API key).")

    model = os.getenv('FIREWORKS_MODEL', DEFAULT_MODEL)
    client = OpenAI(api_key=api_key, base_url=FIREWORKS_BASE_URL, timeout=55, max_retries=1)

    # Truncate very long transcripts to stay within token limits and keep it fast
    max_chars = 80000
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[Transcript truncated for processing]"

    try:
        response = client.chat.completions.create(
            model=model,
            max_tokens=2048,
            temperature=0.2,
            messages=[
                {"role": "system", "content": SUMMARIZE_PROMPT},
                {"role": "user", "content": f"Please summarize this transcript:\n\n{text}"},
            ],
        )
    except AuthenticationError:
        raise SummarizerError("The summarizer is temporarily unavailable (invalid API credentials).")
    except RateLimitError:
        raise SummarizerError("The summarizer is busy right now. Please try again in a moment.")
    except (APIConnectionError, APITimeoutError):
        raise SummarizerError("Could not reach the summarization service. Please try again.")
    except APIStatusError as e:
        if e.status_code in (401, 402, 403, 404):
            raise SummarizerError("The summarizer is temporarily unavailable. Please try again later.")
        raise SummarizerError(f"The summarization service returned an error ({e.status_code}).")

    summary_text = (response.choices[0].message.content or "").strip()
    if not summary_text:
        raise SummarizerError("The summarizer returned an empty result. Please try again.")

    # Build a clean Word document from the summary
    return _build_docx(summary_text)


SECTION_HEADERS = ('OVERVIEW', 'KEY POINTS', 'MAIN DISCUSSION', 'ACTION ITEMS')


def _section_name(line: str):
    """Return the canonical section title if the line is a section header, else None."""
    upper = line.upper().rstrip(':').strip()
    for name in SECTION_HEADERS:
        if upper == name or upper.startswith(name + ' ('):
            return name.title() if name != 'KEY POINTS' else 'Key Points'
    return None


def _build_docx(summary_text: str) -> bytes:
    """Convert the plain-text summary into a formatted .docx file."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = summary_text.strip().split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if line.startswith('TITLE:'):
            title_text = line.replace('TITLE:', '').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(12)

        elif _section_name(line):
            p = doc.add_paragraph()
            run = p.add_run(_section_name(line))
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = 'Calibri'
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)

        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(3)

        elif line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
