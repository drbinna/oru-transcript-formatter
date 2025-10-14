"""
Flask web application with fallback model support.
Tries Claude Sonnet 4.5 first, falls back to 3.5 if needed.
"""

import os
import tempfile
import traceback
import logging
from pathlib import Path
from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename

# Try to import CORS
try:
    from flask_cors import CORS
    CORS_AVAILABLE = True
except ImportError:
    CORS_AVAILABLE = False

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Core dependencies for AI formatting
import anthropic
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Enable CORS if available
if CORS_AVAILABLE:
    CORS(app)
else:
    @app.after_request
    def after_request(response):
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
        return response

app.secret_key = os.environ.get('FLASK_SECRET_KEY', os.urandom(24).hex())
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Global error handlers
@app.errorhandler(404)
def handle_404(e):
    """Return JSON for 404 errors on API endpoints."""
    if request.path.startswith('/upload') or request.path.startswith('/download') or request.path.startswith('/health'):
        return jsonify({'error': 'Not found'}), 404
    return render_template('index.html'), 404

@app.errorhandler(500)
def handle_500(e):
    """Return JSON for 500 errors."""
    logger.error(f"500 error: {str(e)}")
    return jsonify({'error': 'Internal server error', 'details': str(e)}), 500

@app.errorhandler(Exception)
def handle_exception(e):
    """Catch all exceptions and return JSON."""
    logger.error(f"Unhandled exception: {str(e)}", exc_info=True)
    return jsonify({'error': str(e)}), 500

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
ALLOWED_EXTENSIONS = {'txt', 'docx'}

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def format_with_claude_fallback(transcript_text):
    """Format transcript using Claude AI with fallback support."""
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    
    if not api_key:
        raise ValueError("Anthropic API key not found. Please set ANTHROPIC_API_KEY environment variable.")
    
    client = anthropic.Anthropic(api_key=api_key)
    
    system_prompt = """You are an expert transcript formatter. Transform this raw transcript into a polished, professional document.

FORMATTING RULES:
1. SPEAKER NAMES: Bold format - **Dr. Billy Wilson:**
2. SCRIPTURE REFERENCES: Bold format - **1 John 2:18**
3. SONG LYRICS: Italic with music notes - *♪ Amazing grace ♪*
4. Fix all encoding errors (â™ª → ♪, â€™ → ', etc.)
5. Maintain proper paragraph structure
6. Output ONLY the formatted transcript, no explanations"""

    # Try Claude Sonnet 4.5 first
    models_to_try = [
        ("claude-sonnet-4-5-20250929", 8192, "Claude Sonnet 4.5"),
        ("claude-3-5-sonnet-20241022", 8000, "Claude 3.5 Sonnet"),
        ("claude-3-5-sonnet-latest", 8000, "Claude 3.5 Sonnet Latest")
    ]
    
    last_error = None
    for model_id, max_tokens, model_name in models_to_try:
        try:
            logger.info(f"Trying model: {model_id}")
            response = client.messages.create(
                model=model_id,
                max_tokens=max_tokens,
                temperature=0.1,
                system=system_prompt,
                messages=[
                    {
                        "role": "user",
                        "content": f"Please format this transcript:\n\n{transcript_text}"
                    }
                ]
            )
            
            formatted_text = response.content[0].text
            logger.info(f"Successfully used model: {model_name}")
            return formatted_text, model_name
            
        except Exception as e:
            logger.warning(f"Model {model_id} failed: {str(e)}")
            last_error = e
            continue
    
    # If all models fail, raise the last error
    raise RuntimeError(f"All Claude models failed. Last error: {str(last_error)}")

def create_word_document(formatted_text, title, output_path):
    """Create a Word document from formatted text."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import re
        
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Process the formatted text
        paragraphs = formatted_text.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            p = doc.add_paragraph()
            
            # Process formatting
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', para_text)
            
            for part in parts:
                if not part:
                    continue
                
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    # Italic text
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    # Normal text
                    p.add_run(part)
            
            # Set font for the entire paragraph
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        
        doc.save(output_path)
        logger.info(f"Word document saved to: {output_path}")
        
    except ImportError:
        logger.warning("python-docx not available, saving as text file")
        with open(output_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
            f.write(f"Title: {title}\n\n")
            f.write(formatted_text)

@app.route('/')
def index():
    """Main page with upload form."""
    return render_template('index.html')

@app.route('/health')
def health_check():
    """Health check endpoint."""
    return jsonify({'status': 'healthy', 'service': 'transcript-formatter'})

@app.route('/model-test')
def model_test():
    """Test which models are available."""
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return jsonify({'error': 'No API key configured'}), 500
    
    client = anthropic.Anthropic(api_key=api_key)
    results = {}
    
    models = [
        "claude-sonnet-4-5-20250929",
        "claude-3-5-sonnet-20241022",
        "claude-3-5-sonnet-latest"
    ]
    
    for model_id in models:
        try:
            response = client.messages.create(
                model=model_id,
                max_tokens=10,
                messages=[{"role": "user", "content": "Hi"}]
            )
            results[model_id] = "Working"
        except Exception as e:
            results[model_id] = f"Failed: {str(e)}"
    
    return jsonify(results)

@app.route('/upload', methods=['POST', 'OPTIONS'])
def upload_file():
    """Handle file upload and processing."""
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'})
    
    try:
        logger.info("Upload request received")
        
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        if file and allowed_file(file.filename):
            upload_path = None
            try:
                # Save uploaded file
                filename = secure_filename(file.filename)
                upload_path = os.path.join(UPLOAD_FOLDER, filename)
                file.save(upload_path)
                
                # Extract text content
                if filename.lower().endswith('.txt'):
                    with open(upload_path, 'r', encoding='utf-8') as f:
                        transcript_text = f.read()
                elif filename.lower().endswith('.docx'):
                    try:
                        from docx import Document
                        doc = Document(upload_path)
                        transcript_text = '\n'.join([para.text for para in doc.paragraphs])
                    except ImportError:
                        return jsonify({'success': False, 'error': 'DOCX support not available'}), 500
                
                # Format with Claude (with fallback)
                formatted_text, model_used = format_with_claude_fallback(transcript_text)
                
                # Generate output filename
                base_name = os.path.splitext(filename)[0]
                output_filename = f"{base_name}_formatted.docx"
                output_path = os.path.join(OUTPUT_FOLDER, output_filename)
                
                # Export to Word
                create_word_document(formatted_text, f"Formatted: {base_name}", output_path)
                
                # Clean up uploaded file
                if upload_path and os.path.exists(upload_path):
                    os.remove(upload_path)
                
                return jsonify({
                    'success': True, 
                    'filename': output_filename,
                    'download_url': f'/download/{output_filename}',
                    'model_used': model_used
                })
                
            except Exception as e:
                logger.error(f"Processing error: {str(e)}")
                if upload_path and os.path.exists(upload_path):
                    os.remove(upload_path)
                return jsonify({'success': False, 'error': f'Processing failed: {str(e)}'}), 500
        
        return jsonify({'success': False, 'error': 'Invalid file type'}), 400
    
    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Download formatted file."""
    try:
        file_path = os.path.join(OUTPUT_FOLDER, secure_filename(filename))
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return jsonify({'success': False, 'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'success': False, 'error': f'Download failed: {str(e)}'}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    logger.info(f"Starting Flask app on port {port}")
    app.run(host='0.0.0.0', port=port, debug=False)