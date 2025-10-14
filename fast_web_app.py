#!/usr/bin/env python3
"""
Fast-loading ORU Transcript Formatter web app.
"""

from flask import Flask, render_template_string, request, jsonify, send_file
import os
import tempfile
from pathlib import Path

app = Flask(__name__)
app.secret_key = 'oru-transcript-formatter-2025'

# HTML template embedded for faster loading
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>ORU Transcript Formatter</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #003366 0%, #002244 100%);
            color: white;
            margin: 0;
            padding: 20px;
            min-height: 100vh;
        }
        .container {
            max-width: 800px;
            margin: 0 auto;
            text-align: center;
        }
        h1 {
            color: #FFD700;
            font-size: 2.5rem;
            margin-bottom: 1rem;
        }
        .upload-area {
            background: rgba(255, 255, 255, 0.1);
            border: 2px dashed #4A90E2;
            border-radius: 15px;
            padding: 40px;
            margin: 30px 0;
        }
        .btn {
            background: linear-gradient(135deg, #FFD700 0%, #FFC107 100%);
            color: #003366;
            border: none;
            padding: 15px 30px;
            font-size: 1.1rem;
            font-weight: bold;
            border-radius: 25px;
            cursor: pointer;
            margin: 10px;
        }
        .btn:hover {
            transform: translateY(-2px);
        }
        #status {
            margin: 20px 0;
            padding: 15px;
            border-radius: 10px;
            display: none;
        }
        .success { background: #28a745; }
        .error { background: #dc3545; }
        .processing { background: #007bff; }
    </style>
</head>
<body>
    <div class="container">
        <h1>🎓 ORU Transcript Formatter</h1>
        <p style="color: #FFD700; font-size: 1.2rem;">AI-Powered Transcript Formatting • Oral Roberts University</p>
        
        <div class="upload-area">
            <h3>📄 Upload Transcript File</h3>
            <input type="file" id="fileInput" accept=".txt" style="margin: 20px 0;">
            <br>
            <button class="btn" onclick="formatTranscript()">🤖 Format Transcript</button>
        </div>
        
        <div id="status"></div>
        <div id="downloadArea" style="display: none;">
            <a id="downloadLink" class="btn" href="#" download>📥 Download Formatted Document</a>
        </div>
        
        <div style="margin-top: 50px; color: #FFD700;">
            <h3>✨ Features</h3>
            <p>🎯 AI-powered speaker detection • 📖 Scripture reference highlighting • 🎨 Professional formatting</p>
            <p>© 2025 Oral Roberts University • Powered by AI</p>
        </div>
    </div>

    <script>
        function showStatus(message, type) {
            const status = document.getElementById('status');
            status.textContent = message;
            status.className = type;
            status.style.display = 'block';
        }

        function formatTranscript() {
            const fileInput = document.getElementById('fileInput');
            const file = fileInput.files[0];
            
            if (!file) {
                showStatus('Please select a transcript file first.', 'error');
                return;
            }
            
            if (!file.name.endsWith('.txt')) {
                showStatus('Please select a .txt file.', 'error');
                return;
            }
            
            showStatus('🤖 Processing transcript with AI... This may take 1-3 minutes.', 'processing');
            
            const formData = new FormData();
            formData.append('file', file);
            
            fetch('/format', {
                method: 'POST',
                body: formData
            })
            .then(response => response.json())
            .then(data => {
                if (data.success) {
                    showStatus('✅ Transcript formatted successfully!', 'success');
                    document.getElementById('downloadLink').href = '/download/' + data.filename;
                    document.getElementById('downloadArea').style.display = 'block';
                } else {
                    showStatus('❌ Error: ' + data.error, 'error');
                }
            })
            .catch(error => {
                showStatus('❌ Network error: ' + error.message, 'error');
            });
        }
    </script>
</body>
</html>
"""

@app.route('/')
def index():
    """Main page."""
    return render_template_string(HTML_TEMPLATE)

@app.route('/format', methods=['POST'])
def format_transcript():
    """Format transcript endpoint."""
    try:
        # Lazy import to speed up startup
        from transcript_formatter.core.claude_formatter import format_with_claude
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        file = request.files['file']
        if not file:
            return jsonify({'error': 'No file uploaded'})
        
        # Read file
        content = file.read().decode('utf-8')
        if not content.strip():
            return jsonify({'error': 'File is empty'})
        
        # Format with AI
        formatted_text = format_with_claude(content)
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = Path(file.filename).stem.replace('_', ' ').replace('-', ' ')
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(16)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(24)
        
        # Add content
        lines = formatted_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            para = doc.add_paragraph()
            
            if '**' in line:
                parts = re.split(r'(\*\*[^*]+\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    else:
                        para.add_run(part)
            else:
                para.add_run(line)
            
            for run in para.runs:
                run.font.size = Pt(11)
            para.space_after = Pt(6)
        
        # Save document
        output_filename = f"{Path(file.filename).stem}_formatted.docx"
        output_path = os.path.join(tempfile.gettempdir(), output_filename)
        doc.save(output_path)
        
        return jsonify({'success': True, 'filename': output_filename})
        
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/download/<filename>')
def download_file(filename):
    """Download formatted file."""
    try:
        file_path = os.path.join(tempfile.gettempdir(), filename)
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)})

if __name__ == '__main__':
    print("🎓 ORU Transcript Formatter - Fast Startup")
    print("🌐 Starting server at http://localhost:8080")
    print("⚡ Ready for use!")
    
    app.run(debug=False, host='0.0.0.0', port=8080, threaded=True)