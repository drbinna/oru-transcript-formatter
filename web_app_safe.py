"""
Flask web application with safe CORS handling.
Falls back gracefully if flask-cors is not available.
"""

import os
import tempfile
import traceback
import logging
from pathlib import Path
from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename

# Try to import CORS, but don't fail if not available
try:
    from flask_cors import CORS
    CORS_AVAILABLE = True
except ImportError:
    CORS_AVAILABLE = False
    print("WARNING: flask-cors not available, CORS support disabled")

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Core dependencies for AI formatting
import anthropic
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Enable CORS if available
if CORS_AVAILABLE:
    CORS(app)
    logger.info("CORS enabled")
else:
    logger.warning("CORS not enabled - flask-cors not installed")
    # Manual CORS headers as fallback
    @app.after_request
    def after_request(response):
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
        return response

app.secret_key = os.environ.get('FLASK_SECRET_KEY', os.urandom(24).hex())
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure JSON responses
@app.before_request
def ensure_json():
    """Log requests and ensure proper handling."""
    logger.info(f"Request: {request.method} {request.path}")

# Simple test endpoint that always works
@app.route('/api/test')
def api_test():
    """Simple test endpoint that returns JSON."""
    return jsonify({
        'status': 'ok',
        'message': 'API is working',
        'cors_enabled': CORS_AVAILABLE
    })

@app.route('/health')
def health_check():
    """Health check endpoint."""
    return jsonify({
        'status': 'healthy',
        'service': 'transcript-formatter',
        'cors_enabled': CORS_AVAILABLE
    })

# Global error handlers that always return JSON
@app.errorhandler(404)
def handle_404(e):
    """Return JSON for 404 errors."""
    return jsonify({'error': 'Not found', 'path': request.path}), 404

@app.errorhandler(500)
def handle_500(e):
    """Return JSON for 500 errors."""
    logger.error(f"500 error: {str(e)}")
    return jsonify({'error': 'Internal server error'}), 500

@app.errorhandler(Exception)
def handle_exception(e):
    """Catch all exceptions and return JSON."""
    logger.error(f"Unhandled exception: {str(e)}", exc_info=True)
    return jsonify({'error': str(e)}), 500

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
ALLOWED_EXTENSIONS = {'txt', 'docx'}

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Claude AI formatting functionality
def format_with_claude_inline(transcript_text):
    """Format transcript using Claude AI - inline implementation."""
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    
    if not api_key:
        raise ValueError("Anthropic API key not found. Please set ANTHROPIC_API_KEY environment variable.")
    
    client = anthropic.Anthropic(api_key=api_key)
    
    system_prompt = """You are an expert transcript formatter. Transform this raw transcript into a polished, professional document matching broadcast-quality standards.

CRITICAL OUTPUT REQUIREMENTS:
- Output ONLY the formatted transcript content
- NO meta-commentary, section headers, or explanations like "Here's the formatted transcript:"
- Start immediately with the actual transcript text

PRECISE FORMATTING RULES:

1. SPEAKER NAMES - Must be EXACTLY this format:
   **Dr. Billy Wilson:** (bold with colon, followed by space)
   **Female Voice:** (for any female speaker)
   **Congregation:** (for group responses)
   Common speakers: Dr. Billy Wilson, Female Voice, Male Voice, Narrator, Congregation
   
2. SCRIPTURE REFERENCES - Must be bold:
   **1 John 2:18** (no other formatting)
   **Romans 8:28-30** (for verse ranges) 
   **Psalm 23** (when no verse specified)
   Scan thoroughly - Scripture references often appear mid-sentence
   
3. SONG LYRICS - Must be italicized with music notes:
   *♪ Amazing grace, how sweet the sound ♪*
   Each lyric line starts and ends with ♪ symbol
   
4. NUMBERED POINTS in teaching sections:
   Use exact numbers from transcript: "1.", "2.", "3." or "First,", "Second,"
   
5. CHARACTER ENCODING - Fix all encoding errors:
   â™ª → ♪ (music note)
   â€™ → ' (apostrophe)
   â€" → — (em dash)  
   â€œ → " (left quote)
   â€\x9d → " (right quote)
   Ã© → é
   
6. PARAGRAPH STRUCTURE:
   - New paragraph for each speaker change
   - New paragraph for major topic shifts
   - Keep related content together
   - Remove any timestamp markers like [00:00:00]

7. PRESERVE ALL CONTENT:
   - Keep every word, including "um", "uh", repetitions
   - Maintain emotional indicators: (Laughter), (Applause), (Pause)
   - Keep any emphasis or special formatting from original

EXAMPLE OUTPUT FORMAT:
**Dr. Billy Wilson:** Welcome to Chapel today. Today we're looking at **Romans 8:28**, which tells us that all things work together for good.

*♪ Blessed assurance, Jesus is mine ♪*

**Female Voice:** That's a powerful truth we need to remember.

**Congregation:** Amen!

Remember: Output ONLY the formatted transcript, nothing else."""

    user_prompt = f"Format this transcript according to the requirements:\n\n{transcript_text}"
    
    try:
        response = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=8000,
            temperature=0.2,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        
        formatted_text = response.content[0].text
        
        # Clean up any potential wrapper text
        lines_to_skip = [
            "Here's the formatted transcript:",
            "Here is the formatted transcript:",
            "---",
            "___"
        ]
        
        lines = formatted_text.split('\n')
        while lines and any(skip in lines[0] for skip in lines_to_skip):
            lines.pop(0)
        
        return '\n'.join(lines).strip()
        
    except Exception as e:
        logger.error(f"Claude API error: {str(e)}")
        raise

# Word export functionality
def export_to_word(formatted_text, title, output_path):
    """Export formatted text to a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Process the formatted text
        paragraphs = formatted_text.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            p = doc.add_paragraph()
            
            # Process the paragraph text for formatting
            import re
            
            # Pattern to match bold text and italic text
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', para_text)
            
            for part in parts:
                if not part:
                    continue
                
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    # Italic text
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    # Normal text
                    p.add_run(part)
            
            # Set font for the entire paragraph
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        
        doc.save(output_path)
        logger.info(f"Word document saved to: {output_path}")
        
    except ImportError:
        logger.warning("python-docx not available, saving as text file")
        # Fallback to text file if python-docx not available
        with open(output_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
            f.write(f"Title: {title}\n\n")
            f.write(formatted_text)

@app.route('/')
def index():
    """Main page with upload form."""
    return render_template('index.html')

@app.route('/upload', methods=['POST', 'OPTIONS'])
def upload_file():
    """Handle file upload and processing."""
    # Handle preflight OPTIONS request
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'})
    
    try:
        logger.info("=== UPLOAD REQUEST START ===")
        
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        if file and allowed_file(file.filename):
            upload_path = None
            try:
                # Save uploaded file
                filename = secure_filename(file.filename)
                upload_path = os.path.join(UPLOAD_FOLDER, filename)
                file.save(upload_path)
                
                # Extract text content
                if filename.lower().endswith('.txt'):
                    with open(upload_path, 'r', encoding='utf-8') as f:
                        transcript_text = f.read()
                elif filename.lower().endswith('.docx'):
                    try:
                        from docx import Document
                        doc = Document(upload_path)
                        transcript_text = '\n'.join([para.text for para in doc.paragraphs])
                    except ImportError:
                        return jsonify({'success': False, 'error': 'DOCX support not available'}), 500
                
                # Format with Claude
                formatted_text = format_with_claude_inline(transcript_text)
                
                # Generate output filename
                base_name = os.path.splitext(filename)[0]
                output_filename = f"{base_name}_formatted.docx"
                output_path = os.path.join(OUTPUT_FOLDER, output_filename)
                
                # Export to Word
                export_to_word(formatted_text, f"Formatted: {base_name}", output_path)
                
                # Clean up uploaded file
                if upload_path and os.path.exists(upload_path):
                    os.remove(upload_path)
                
                return jsonify({
                    'success': True, 
                    'filename': output_filename,
                    'download_url': f'/download/{output_filename}'
                })
                
            except Exception as e:
                logger.error(f"Processing error: {str(e)}")
                if upload_path and os.path.exists(upload_path):
                    os.remove(upload_path)
                return jsonify({'success': False, 'error': f'Processing failed: {str(e)}'}), 500
        
        return jsonify({'success': False, 'error': 'Invalid file type'}), 400
    
    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Download formatted file."""
    try:
        file_path = os.path.join(OUTPUT_FOLDER, secure_filename(filename))
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return jsonify({'success': False, 'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'success': False, 'error': f'Download failed: {str(e)}'}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    logger.info(f"Starting Flask app on port {port}")
    logger.info(f"CORS enabled: {CORS_AVAILABLE}")
    app.run(host='0.0.0.0', port=port, debug=False)