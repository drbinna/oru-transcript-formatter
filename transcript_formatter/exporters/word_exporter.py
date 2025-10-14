from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

class WordExporter:
    def __init__(self):
        """Initialize Word exporter with default professional styling"""
        self.doc = Document()
        self._setup_styles()
        self._setup_page_layout()
    
    def _setup_page_layout(self):
        """Set up professional page margins"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
    
    def _setup_styles(self):
        """Set up document styles"""
        styles = self.doc.styles
        
        # Normal/Body text style
        normal_style = styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
        
        # Heading 1 style (for numbered sections)
        try:
            h1_style = styles['Heading 1']
            h1_style.font.name = 'Calibri'
            h1_style.font.size = Pt(14)
            h1_style.font.bold = True
            h1_style.paragraph_format.space_before = Pt(12)
            h1_style.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass  # Style doesn't exist, skip
    
    def _add_horizontal_line(self):
        """Add a professional horizontal line"""
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        
        # Create border element
        pBdr = OxmlElement('w:pBdr')
        
        # Add bottom border
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # 0.75pt thickness
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')  # Black
        
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        return p
    
    def export(self, formatted_text, output_path):
        """Export formatted markdown to Word document"""
        
        lines = formatted_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines but maintain spacing
            if not line:
                self.doc.add_paragraph()
                i += 1
                continue
            
            # Detect divider lines (optional)
            if re.match(r'^[_\-]{10,}', line):
                self._add_horizontal_line()
                i += 1
                continue
            
            # Title - first bold text (centered)
            if line.startswith('**') and line.endswith('**') and len(self.doc.paragraphs) < 5:
                title = line.strip('*')
                p = self.doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                run = p.add_run(title)
                run.font.name = 'Arial'
                run.font.size = Pt(20)
                run.bold = True
                
                p.paragraph_format.space_after = Pt(12)
                i += 1
                continue
            
            # Numbered teaching headers (1. Title, 2. Title, etc.)
            if re.match(r'\*\*\d+\.\s+[^*]+\*\*', line):
                header_text = line.strip('*')
                p = self.doc.add_heading(header_text, level=1)
                i += 1
                continue
            
            # Song lyrics
            if line.startswith('♪'):
                lyric_text = line.rstrip('♪').strip()
                p = self.doc.add_paragraph()
                
                # Standalone music notes
                if lyric_text == '♪♪♪' or lyric_text == '♪ ♪ ♪':
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run(lyric_text)
                else:
                    # Regular lyrics - left aligned, italic, gray
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    run = p.add_run(lyric_text)
                    run.italic = True
                    run.font.color.rgb = RGBColor(89, 89, 89)
                
                p.paragraph_format.space_after = Pt(0)
                i += 1
                continue
            
            # Regular paragraph with inline formatting
            p = self.doc.add_paragraph()
            self._add_formatted_text(p, line)
            i += 1
        
        self.doc.save(output_path)
        return output_path
    
    def _add_formatted_text(self, paragraph, text):
        """Add text with proper formatting - NO asterisks in output"""
        
        # Find all bold sections
        bold_pattern = r'\*\*(.+?)\*\*'
        bold_matches = list(re.finditer(bold_pattern, text))
        
        # Find all italic sections
        italic_pattern = r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)'
        italic_matches = list(re.finditer(italic_pattern, text))
        
        # Build formatting regions
        regions = []
        for match in bold_matches:
            regions.append({
                'type': 'bold',
                'start': match.start(),
                'end': match.end(),
                'content': match.group(1)
            })
        
        for match in italic_matches:
            # Skip if inside bold
            inside_bold = any(
                r['start'] < match.start() < r['end'] 
                for r in regions if r['type'] == 'bold'
            )
            if not inside_bold:
                regions.append({
                    'type': 'italic',
                    'start': match.start(),
                    'end': match.end(),
                    'content': match.group(1)
                })
        
        regions.sort(key=lambda x: x['start'])
        
        # Build paragraph
        last_pos = 0
        
        for region in regions:
            # Add text before formatted region
            if region['start'] > last_pos:
                before_text = text[last_pos:region['start']].replace('*', '')
                if before_text:
                    paragraph.add_run(before_text)
            
            # Add formatted content
            run = paragraph.add_run(region['content'])
            
            if region['type'] == 'bold':
                run.bold = True
                # Scripture references in blue
                if self._is_scripture_reference(region['content']):
                    run.font.color.rgb = RGBColor(5, 99, 193)
            
            elif region['type'] == 'italic':
                run.italic = True
                # Long quotes in gray
                if len(region['content']) > 50:
                    run.font.color.rgb = RGBColor(89, 89, 89)
            
            last_pos = region['end']
        
        # Add remaining text
        if last_pos < len(text):
            remaining = text[last_pos:].replace('*', '')
            if remaining:
                paragraph.add_run(remaining)
    
    def _is_scripture_reference(self, text):
        """Check if text is a scripture reference"""
        # Has verse numbers
        if re.search(r'\d+:\d+', text) or re.search(r'\d+--\d+', text):
            return True
        
        # Common book names
        books = [
            'John', 'Timothy', 'Mark', 'Jeremiah', 'Hebrews',
            'Luke', 'Acts', 'Jude', 'Matthew', 'Romans',
            'Corinthians', 'Genesis', 'Exodus', 'Psalms'
        ]
        return any(book in text for book in books)