"""
ORU Transcript Formatter - Hugging Face Spaces Deployment
AI-Powered Transcript Formatting with ORU Branding
"""

import os
import tempfile
from pathlib import Path
import gradio as gr
import sys

# Add current directory to Python path for imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

try:
    from transcript_formatter.core.claude_formatter import format_with_claude
except ImportError:
    # Fallback for Hugging Face deployment
    from dotenv import load_dotenv
    import anthropic
    
    load_dotenv()
    
    def format_with_claude(text):
        """Simplified Claude formatter for Hugging Face."""
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY not found in environment variables")
        
        client = anthropic.Anthropic(api_key=api_key)
        
        system_prompt = """You are a professional transcript formatter. Format this transcript with:
        1. Bold speaker names: **Speaker Name:**
        2. Bold Scripture references: **1 John 2:18**, **Mark chapter 13 verse 13**
        3. Fix character encoding issues
        4. Create proper paragraph breaks
        5. Remove timestamps
        6. Return clean markdown format
        Preserve all original content and meaning."""
        
        message = client.messages.create(
            model="claude-3-5-sonnet-20240620",
            max_tokens=8000,
            temperature=0.1,
            system=system_prompt,
            messages=[{"role": "user", "content": f"Please format this transcript:\n\n{text}"}]
        )
        
        return message.content[0].text

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_word_document(formatted_text, title):
    """Create a Word document from formatted text."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(24)
    
    # Add metadata
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run("Formatted with AI • ORU Transcript Formatter")
    meta_run.font.size = Pt(10)
    meta_run.italic = True
    meta_para.space_after = Pt(12)
    
    # Add separator
    doc.add_paragraph("_" * 50).space_after = Pt(12)
    
    # Process the formatted text and add to document
    lines = formatted_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        para = doc.add_paragraph()
        
        # Check if line contains bold formatting (markdown style)
        if '**' in line:
            # Parse markdown-style bold formatting
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    bold_text = part[2:-2]  # Remove ** markers
                    run = para.add_run(bold_text)
                    run.bold = True
                else:
                    # Regular text
                    para.add_run(part)
        else:
            # Regular paragraph
            para.add_run(line)
        
        # Set font size
        for run in para.runs:
            run.font.size = Pt(11)
        
        para.space_after = Pt(6)
    
    return doc

def format_transcript(file):
    """Format a transcript file using AI."""
    if file is None:
        return None, "Please upload a transcript file."
    
    try:
        # Read the uploaded file
        if file.name.endswith('.txt'):
            with open(file.name, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            return None, "Please upload a .txt file."
        
        # Format using AI
        formatted_text = format_with_claude(content)
        
        # Create Word document
        title = Path(file.name).stem.replace('_', ' ').replace('-', ' ')
        doc = create_word_document(formatted_text, title)
        
        # Save to temporary file
        output_path = tempfile.mktemp(suffix='.docx')
        doc.save(output_path)
        
        # Return file and success message
        return output_path, "✅ Transcript formatted successfully! Download your Word document below."
        
    except Exception as e:
        return None, f"❌ Error formatting transcript: {str(e)}"

# Custom CSS for ORU branding
css = """
.gradio-container {
    background: linear-gradient(135deg, #003366 0%, #002244 100%) !important;
    color: white !important;
}

.gr-button-primary {
    background: linear-gradient(135deg, #FFD700 0%, #FFC107 100%) !important;
    color: #003366 !important;
    border: none !important;
    font-weight: bold !important;
}

.gr-button-primary:hover {
    background: linear-gradient(135deg, #FFC107 0%, #FFB300 100%) !important;
    transform: translateY(-1px) !important;
}

h1 {
    color: #FFD700 !important;
    text-align: center !important;
    font-size: 2.5rem !important;
    margin-bottom: 1rem !important;
}

.gr-form {
    background: rgba(255, 255, 255, 0.1) !important;
    border-radius: 15px !important;
    padding: 2rem !important;
    backdrop-filter: blur(10px) !important;
}

.gr-file {
    border: 2px dashed #4A90E2 !important;
    border-radius: 10px !important;
    background: rgba(255, 255, 255, 0.05) !important;
}

.footer {
    text-align: center !important;
    color: #FFD700 !important;
    margin-top: 2rem !important;
}
"""

# Create Gradio interface
with gr.Blocks(css=css, title="ORU Transcript Formatter") as demo:
    gr.HTML("""
    <h1>🎓 ORU Transcript Formatter</h1>
    <p style="text-align: center; color: #FFD700; font-size: 1.2rem; margin-bottom: 2rem;">
        AI-Powered Transcript Formatting • Oral Roberts University
    </p>
    """)
    
    with gr.Row():
        with gr.Column():
            file_input = gr.File(
                label="📄 Upload Transcript File (.txt)",
                file_types=[".txt"],
                type="filepath"
            )
            
            format_btn = gr.Button(
                "🤖 Format Transcript",
                variant="primary",
                size="lg"
            )
            
            status_output = gr.Textbox(
                label="Status",
                interactive=False,
                lines=2
            )
        
        with gr.Column():
            file_output = gr.File(
                label="📥 Download Formatted Document",
                interactive=False
            )
    
    gr.HTML("""
    <div class="footer">
        <h3>✨ Features</h3>
        <p>🎯 AI-powered speaker detection • 📖 Scripture reference highlighting • 🎨 Professional formatting</p>
        <p>© 2025 Oral Roberts University • Powered by AI</p>
    </div>
    """)
    
    # Connect the interface
    format_btn.click(
        fn=format_transcript,
        inputs=[file_input],
        outputs=[file_output, status_output]
    )

# Launch the demo
if __name__ == "__main__":
    demo.launch(
        server_name="0.0.0.0",
        server_port=7860,
        share=False
    )